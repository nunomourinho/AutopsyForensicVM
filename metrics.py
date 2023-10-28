import pandas as pd
from scipy.stats import pointbiserialr
from statsmodels.formula.api import ols
import statsmodels.api as sm
from docx import Document

def analyze_virtualization_metrics_final(file_path):
    # 1. Read the Excel File
    data = pd.read_excel(file_path)
    
    # 2. Data Analysis
    # Descriptive statistics for conversion time based on tool and conversion mode
    descriptive_stats = data.groupby(['Tool used to mount image', 'Conversion mode'])['Real conversion time in seconds'].describe()

    # Point-biserial correlation between tool used and conversion time
    data['Tool_Binarized'] = data['Tool used to mount image'].apply(lambda x: 0 if x == 'aff' else 1)
    correlation, p_value = pointbiserialr(data['Tool_Binarized'], data['Real conversion time in seconds'])
    
    # ANOVA for conversion time based on tool and conversion mode
    data_renamed = data.rename(columns={
        'Real conversion time in seconds': 'ConversionTime',
        'Tool used to mount image': 'Tool',
        'Conversion mode': 'Mode'
    })
    model_renamed = ols('ConversionTime ~ C(Tool) * C(Mode)', data=data_renamed).fit()
    anova_table_renamed = sm.stats.anova_lm(model_renamed, typ=2)
    
    # 3. Generate Word Document
    doc = Document()
    doc.add_heading('Virtualization Metrics Analysis', 0)
    
    # Descriptive Statistics
    doc.add_heading('Descriptive Statistics', level=1)
    table = doc.add_table(rows=1, cols=9)
    hdr_cells = table.rows[0].cells
    columns = ['Tool', 'Mode', 'count', 'mean', 'std', 'min', '25%', '50%', 'max']
    for i, col in enumerate(columns):
        hdr_cells[i].text = col.capitalize()

    for index, row in descriptive_stats.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(index[0])
        cells[1].text = str(index[1])
        for i in range(2, 9):
            cells[i].text = str(round(row[columns[i]], 2))
    
    # Point-biserial Correlation
    doc.add_heading('Point-biserial Correlation', level=1)
    doc.add_paragraph(f'Correlation: {correlation:.3f}')
    doc.add_paragraph(f'p-value: {p_value:.3f}')
    
    # ANOVA Results
    doc.add_heading('ANOVA Results', level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    columns = ['Source', 'sum_sq', 'df', 'F']
    for i, col in enumerate(columns):
        hdr_cells[i].text = col.capitalize()

    for index, row in anova_table_renamed.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(index)
        for i in range(1, 4):
            cells[i].text = str(round(row[columns[i]], 2))
    
    # Save the document
    doc_path = "virtualization_metrics_analysis_final.docx"
    doc.save(doc_path)
    
    # 4. Generate LaTeX Document    
    latex_content = "\\section*{Virtualization Metrics Analysis}\n"
    latex_content += "\\subsection*{Descriptive Statistics}\n"
    latex_content += descriptive_stats.to_latex()
    latex_content += "\n\\subsection*{Point-biserial Correlation}\n"
    latex_content += f"Correlation: {correlation:.3f}\n\np-value: {p_value:.3f}\n"
    latex_content += "\n\\subsection*{ANOVA Results}\n"
    latex_content += anova_table_renamed.to_latex()    
    
    # Save the LaTeX content to a .tex file
    latex_path = "virtualization_metrics_analysis_final.tex"
    with open(latex_path, 'w') as latex_file:
        latex_file.write(latex_content)
    
    return doc_path, latex_path

# Execute the function and get the paths to the generated documents
generated_doc_path_final, generated_latex_path_final = analyze_virtualization_metrics_final("virtualization_metrics.xls")

